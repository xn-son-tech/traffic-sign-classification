import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_proposal():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("PROJECT PROPOSAL\n")
    run_title.font.size = Pt(16)
    run_title.bold = True
    
    run_sub = title.add_run("Traffic Sign Classification Using PCA and Support Vector Machine (SVM)")
    run_sub.font.size = Pt(14)
    run_sub.bold = True
    
    doc.add_paragraph("\n")
    
    # 1. Introduction
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. Introduction")
    r1.font.size = Pt(13)
    r1.bold = True
    
    p1_1 = doc.add_paragraph()
    p1_1.add_run(
        "Traffic sign recognition (TSR) is an essential component of Intelligent Transportation Systems (ITS) "
        "and Advanced Driver Assistance Systems (ADAS). The ability to accurately and quickly recognize traffic signs "
        "improves road safety, supports autonomous driving decision-making, and reduces human errors in traffic monitoring. "
        "While deep learning approaches have achieved impressive results in computer vision tasks, they typically demand "
        "significant computational resources (such as GPUs) and lack mathematical interpretability. In contrast, traditional "
        "Machine Learning techniques remain highly valuable due to their lower computational footprints, ease of deployment "
        "on embedded vehicle systems, and solid mathematical foundations. This project focuses on utilizing a classic yet "
        "powerful combination: Principal Component Analysis (PCA) for dimensionality reduction and Support Vector Machine (SVM) "
        "for multi-class classification, applied to the German Traffic Sign Recognition Benchmark (GTSRB) dataset."
    )
    
    h1_sub1 = doc.add_paragraph()
    r1_sub1 = h1_sub1.add_run("Machine Learning Problem:")
    r1_sub1.bold = True
    p1_2 = doc.add_paragraph()
    p1_2.add_run(
        "The problem presented in this project is a supervised, multi-class image classification task. "
        "Given an input image containing a single traffic sign, the goal is to map the image features to one of the 43 "
        "predefined traffic sign categories (e.g., speed limits, stop signs, yield signs, danger warnings). Formally, "
        "for an input image vector x, the model must predict a class label y \u2208 {0, 1, ..., 42}."
    )
    
    h1_sub2 = doc.add_paragraph()
    r1_sub2 = h1_sub2.add_run("Contributions & Objectives:")
    r1_sub2.bold = True
    
    contribs = [
        "Develop an end-to-end Machine Learning pipeline that processes high-dimensional traffic sign images, compresses them, and classifies them using PCA + SVM.",
        "Perform statistical analysis and exploratory data analysis (EDA) on the GTSRB dataset to understand class distributions and image properties.",
        "Implement data visualization techniques (PCA, t-SNE/UMAP) to project the high-dimensional image space into 2D/3D spaces to visually examine class separability.",
        "Analyze the impact of PCA components on classification accuracy, training time, and inference latency, establishing a trade-off curve between speed and accuracy.",
        "Provide a comprehensive evaluation of the model using standard metrics (Accuracy, Precision, Recall, F1-Score, Confusion Matrix) and complexity analysis."
    ]
    for c in contribs:
        doc.add_paragraph(c, style='List Bullet')
        
    doc.add_paragraph("\n")
    
    # 2. Project Process/Methodology
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. Project Process / Methodology")
    r2.font.size = Pt(13)
    r2.bold = True
    
    p2_1 = doc.add_paragraph()
    p2_1.add_run(
        "The project methodology follows a structured pipeline designed to ingest, clean, reduce, and classify "
        "traffic sign data efficiently. The main architecture consists of four distinct phases: Data Preprocessing, "
        "Data Analysis & Visualization, Dimensionality Reduction via PCA, and SVM Classification."
    )
    
    h2_sub1 = doc.add_paragraph()
    r2_sub1 = h2_sub1.add_run("Input and Expected Output:")
    r2_sub1.bold = True
    p2_io = doc.add_paragraph()
    p2_io.add_run(
        "Input: Raw, variable-sized RGB images of traffic signs from the GTSRB dataset.\n"
        "Pipeline: Raw Image \u2192 Resizing (32x32) \u2192 Grayscale & Normalization \u2192 Flattening \u2192 PCA projection \u2192 SVM Predictor \u2192 Output Label.\n"
        "Output: Predicted class label y \u2208 {0, 1, ..., 42} corresponding to the traffic sign category."
    )
    
    h2_sub2 = doc.add_paragraph()
    r2_sub2 = h2_sub2.add_run("Data Analysis:")
    r2_sub2.bold = True
    p2_da = doc.add_paragraph()
    p2_da.add_run(
        "We will perform exploratory statistical analysis on the German Traffic Sign Recognition Benchmark (GTSRB). "
        "The dataset contains over 50,000 physical images of traffic signs across 43 classes. We will study:\n"
        "1. Class Imbalance: Analyzing the distribution of samples per class. Some classes have less than 200 samples, "
        "while others have more than 2,000. This statistical analysis is critical to decide if class weighting or "
        "stratified splitting is required.\n"
        "2. Image Dimensions: The original images have varying resolutions, ranging from 15x15 to 250x250 pixels. Statistical "
        "properties such as mean, median, and standard deviation of width/height will be calculated."
    )
    
    h2_sub3 = doc.add_paragraph()
    r2_sub3 = h2_sub3.add_run("Data Preprocessing:")
    r2_sub3.bold = True
    p2_prep = doc.add_paragraph()
    p2_prep.add_run(
        "To ensure consistent inputs for PCA and SVM, the following preprocessing steps will be executed:\n"
        "- Image Resizing: All images will be resized to a uniform dimension of 32x32 pixels to maintain spatial balance and reasonable features.\n"
        "- Color Space Conversion: Images will be converted from RGB to Grayscale to reduce input features from 3072 (32x32x3) to 1024 (32x32x1), "
        "eliminating variations caused by lighting color temperatures.\n"
        "- Contrast Enhancement: Applying Contrast Limited Adaptive Histogram Equalization (CLAHE) to handle severe shadows, "
        "over-exposure, and low-contrast conditions typical in real-world driving scenes.\n"
        "- Normalization: Normalizing pixel intensities to [0, 1] range to avoid numerical instability in downstream algorithms.\n"
        "- Missing/Noisy Data Handling: Filtering out highly corrupted or extremely low-resolution images that are unreadable."
    )
    
    h2_sub4 = doc.add_paragraph()
    r2_sub4 = h2_sub4.add_run("Data Visualization in Low-Dimensional Space:")
    r2_sub4.bold = True
    p2_vis = doc.add_paragraph()
    p2_vis.add_run(
        "High-dimensional image data (even after flattening to 1024 features) is impossible to perceive directly. "
        "We will implement dimensionality reduction methods specifically for 2D and 3D data visualization:\n"
        "- Principal Component Analysis (PCA): Projecting the data onto the first 2 or 3 principal components to observe the main axes of variance.\n"
        "- t-Distributed Stochastic Neighbor Embedding (t-SNE) / UMAP: Non-linear visualization methods to map the high-dimensional clusters "
        "into 2D space. This will allow us to visually verify whether different classes of traffic signs (e.g., circular warning signs vs. triangular danger signs) "
        "form distinct, separable clusters before feeding them into the SVM."
    )
    
    doc.add_paragraph("\n")
    
    # 3. Experiments
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. Experiments")
    r3.font.size = Pt(13)
    r3.bold = True
    
    h3_sub1 = doc.add_paragraph()
    r3_sub1 = h3_sub1.add_run("Experimental Protocols & Model Parameters:")
    r3_sub1.bold = True
    p3_proto = doc.add_paragraph()
    p3_proto.add_run(
        "To optimize the PCA + SVM pipeline, we will conduct structured grid search experiments:\n"
        "- PCA Components (k): We will experiment with different numbers of principal components (e.g., k = 30, 50, 100, 200, 300) "
        "to determine the elbow point where cumulative explained variance exceeds 90% or 95%.\n"
        "- SVM Hyperparameters: Support Vector Machines are sensitive to hyperparameters. We will perform Grid Search Cross-Validation "
        "to tune:\n"
        "  * Kernel Type: Linear vs. Radial Basis Function (RBF) to check if non-linear boundaries significantly improve accuracy.\n"
        "  * Regularization Parameter (C): Testing C \u2208 {0.1, 1, 10, 100} to balance margin size and training classification errors.\n"
        "  * Kernel Coefficient (gamma): For RBF kernel, testing gamma \u2208 {'scale', 'auto', 0.001, 0.01}."
    )
    
    h3_sub2 = doc.add_paragraph()
    r3_sub2 = h3_sub2.add_run("Dataset Split & Setup:")
    r3_sub2.bold = True
    p3_setup = doc.add_paragraph()
    p3_setup.add_run(
        "The GTSRB dataset will be partitioned into training and testing subsets using a stratified split (e.g., 80% training, 20% testing). "
        "Stratification ensures that the high imbalance in traffic sign classes is proportionally preserved in both subsets, "
        "preventing models from performing poorly on underrepresented classes."
    )
    
    h3_sub3 = doc.add_paragraph()
    r3_sub3 = h3_sub3.add_run("Computational Complexity & Running Time:")
    r3_sub3.bold = True
    p3_comp = doc.add_paragraph()
    p3_comp.add_run(
        "A critical part of the experiment is analyzing complexity:\n"
        "- Training Complexity: Training standard SVM is O(d * N^2) to O(d * N^3), where N is the number of training samples and d is the feature size. "
        "With N \u2248 39,000, running SVM on raw 1024-dimensional space is extremely slow. By using PCA to reduce d from 1024 to k \u2248 100, "
        "we will mathematically and empirically demonstrate massive speedups in training time.\n"
        "- Empirical Run Time: We will record and plot: (1) Total pipeline training time, and (2) Average inference latency per image (target: < 10ms for real-time ADAS compatibility) "
        "across varying PCA components."
    )
    
    h3_sub4 = doc.add_paragraph()
    r3_sub4 = h3_sub4.add_run("Evaluation & Discussion:")
    r3_sub4.bold = True
    p3_eval = doc.add_paragraph()
    p3_eval.add_run(
        "The classification performance will be comprehensively evaluated using:\n"
        "- Accuracy: Overall correct predictions.\n"
        "- Macro and Weighted Precision, Recall, and F1-Score: Crucial due to class imbalances in GTSRB.\n"
        "- Confusion Matrix: To pinpoint exactly which classes are frequently misclassified (e.g., misclassifying 30km/h and 80km/h speed limit signs due to visual similarity).\n"
        "In the discussion phase, we will analyze the trade-offs of reducing features via PCA. We will answer: "
        "How much information (variance) is discarded, how does it affect classification boundaries, and how can we optimize the pipeline for deployment in resource-constrained systems?"
    )
    
    doc.add_paragraph("\n")
    
    # 4. Conclusions and perspectives
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. Conclusions and perspectives")
    r4.font.size = Pt(13)
    r4.bold = True
    
    h4_sub1 = doc.add_paragraph()
    r4_sub1 = h4_sub1.add_run("Completed Work / Conclusion:")
    r4_sub1.bold = True
    p4_concl = doc.add_paragraph()
    p4_concl.add_run(
        "This project will establish a highly optimized, mathematically grounded PCA + SVM pipeline for classifying traffic signs from the GTSRB dataset. "
        "Through data preprocessing, structured visual exploration, hyperparameter tuning, and detailed complexity analysis, we expect to demonstrate "
        "that traditional Machine Learning algorithms, when designed correctly, can achieve highly competitive accuracy and outstanding inference speeds "
        "without the need for massive Deep Learning infrastructures."
    )
    
    h4_sub2 = doc.add_paragraph()
    r4_sub2 = h4_sub2.add_run("Future Work / Perspectives:")
    r4_sub2.bold = True
    p4_future = doc.add_paragraph()
    p4_future.add_run(
        "For future iterations and extension, the following perspectives are proposed:\n"
        "1. Real-time Object Detection: Integrating a localization network (e.g., Sliding Window or light Haar Cascades) to detect and locate signs before classifying them.\n"
        "2. Robustness to Adverse Weather: Testing and enhancing the system under simulated adverse weather conditions (heavy rain, fog, nighttime glare) using advanced noise filtering or data augmentation.\n"
        "3. Comparison with Deep Learning: Eventually benchmarking this PCA+SVM model against lightweight Convolutional Neural Networks (like MobileNet or SqueezeNet) to evaluate the tradeoff between micro-second CPU latency and deep-learning precision."
    )
    
    doc.add_paragraph("\n")
    
    # Team Division
    h5 = doc.add_paragraph()
    r5 = h5.add_run("5. Team Responsibilities")
    r5.font.size = Pt(13)
    r5.bold = True
    
    p5_team = doc.add_paragraph()
    p5_team.add_run(
        "Member A: Data Engineering & Visualization\n"
        "- Literature review & Dataset acquisition.\n"
        "- Exploratory Data Analysis (EDA) & preprocessing pipeline (resizing, histogram equalization, scaling).\n"
        "- High-dimensional data visualization using PCA, t-SNE, and UMAP.\n"
        "- Report writing & documentation.\n\n"
        "Member B: Model Development, Optimization & Evaluation\n"
        "- Implementing PCA dimensionality reduction.\n"
        "- Developing the SVM classifier & tuning hyperparameters (C, gamma, kernel) via Grid Search CV.\n"
        "- Analyzing training/inference time and computational complexity (Big-O).\n"
        "- Evaluating model using Accuracy, Precision, Recall, F1, and Confusion Matrix."
    )
    
    doc.save(r'd:\school\master of engineering\S2\MLE501.22-machine-learning\project\proposal.docx')
    print("New proposal.docx created successfully!")

if __name__ == '__main__':
    create_proposal()
