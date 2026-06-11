import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    """Set the background color of a table cell (fill_color in hex format)."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_report():
    doc = docx.Document()
    
    # 1. Page Setup (Margins: 1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # 2. Styles Setup (Times New Roman, 12pt default)
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(51, 51, 51) # Dark charcoal instead of pure black for premium feel
    
    # Custom Heading Function to apply uniform styling
    def add_heading(text, level, space_before=12, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        if level == 1:
            run = p.add_run(text)
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(47, 85, 151) # Accent Navy
        elif level == 2:
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(47, 85, 151)
        elif level == 3:
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(47, 85, 151)
        return p

    def add_paragraph(text="", space_after=6, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if text:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = bold
            run.font.italic = italic
        return p

    def add_bullet(text, space_after=4):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    def add_numbered_bullet(text, num_str, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        run_num = p.add_run(num_str + " ")
        run_num.font.name = 'Times New Roman'
        run_num.font.bold = True
        run_num.font.size = Pt(12)
        
        run_text = p.add_run(text)
        run_text.font.name = 'Times New Roman'
        run_text.font.size = Pt(12)
        return p

    def add_math_block(eq_text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(eq_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True
        return p

    def add_code_block(code_text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(5.8)
        cell = table.cell(0, 0)
        set_cell_background(cell, "F5F5F5")
        
        # Apply border style (left border only)
        # For simplicity, we just use shading, which looks extremely clean.
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 30, 30)
        
        # Add space after table
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_image(img_name, caption, width_in=5.0):
        brain_dir = r"C:\Users\ACER\.gemini\antigravity-ide\brain\e43fa4ef-2890-4f04-a9c0-41a1d14c0a00"
        img_path = os.path.join(brain_dir, img_name)
        if not os.path.exists(img_path):
            print(f"Warning: Image {img_name} not found at {img_path}!")
            return
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_in))
        
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(12)
        caption_p.paragraph_format.keep_with_next = True
        
        c_run = caption_p.add_run(caption)
        c_run.font.name = 'Times New Roman'
        c_run.font.size = Pt(10)
        c_run.font.italic = True
        c_run.font.color.rgb = RGBColor(100, 100, 100)

    def add_styled_table(headers, data, alignments=None):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Format headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_background(hdr_cells[i], "2F5597") # Accent Navy
            p = hdr_cells[i].paragraphs[0]
            p.alignment = alignments[i] if alignments else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                
        # Format data rows
        for row_idx, row_data in enumerate(data):
            row_cells = table.add_row().cells
            for col_idx, text in enumerate(row_data):
                row_cells[col_idx].text = str(text)
                p = row_cells[col_idx].paragraphs[0]
                p.alignment = alignments[col_idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                # Alternating row colors
                if row_idx % 2 == 1:
                    set_cell_background(row_cells[col_idx], "F2F2F2")
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_callout(text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(5.8)
        cell = table.cell(0, 0)
        set_cell_background(cell, "F2F5F8")
        
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        run_lbl = p.add_run("NOTE: ")
        run_lbl.font.name = 'Times New Roman'
        run_lbl.font.bold = True
        run_lbl.font.size = Pt(10.5)
        run_lbl.font.italic = True
        run_lbl.font.color.rgb = RGBColor(47, 85, 151)
        
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run.font.italic = True
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # DOCUMENT GENERATION
    # ----------------------------------------------------
    
    # Title Page
    title_p = add_paragraph(align=WD_ALIGN_PARAGRAPH.CENTER)
    title_p.paragraph_format.space_before = Pt(72)
    title_run = title_p.add_run("Traffic Sign Classification Using Principal Component Analysis\nand Support Vector Machine\n")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(47, 85, 151)
    
    meta_p = add_paragraph(align=WD_ALIGN_PARAGRAPH.CENTER)
    meta_p.paragraph_format.space_before = Pt(24)
    r_meta = meta_p.add_run(
        "Course: MLE501.22 – Machine Learning\n"
        "Project Type: Final Project Report\n"
        "Dataset: German Traffic Sign Recognition Benchmark (GTSRB)"
    )
    r_meta.font.size = Pt(12)
    r_meta.font.italic = True
    
    doc.add_page_break()
    
    # Abstract
    add_heading("Abstract", level=2, space_before=18, space_after=12)
    p_abstract = add_paragraph(
        "This report presents a complete, end-to-end machine learning pipeline for the multi-class classification of traffic signs from the German Traffic Sign Recognition Benchmark (GTSRB) dataset. The proposed approach integrates three core components: (1) an image preprocessing pipeline incorporating Region-of-Interest (ROI) cropping, grayscale conversion, and Contrast Limited Adaptive Histogram Equalization (CLAHE) to produce normalized, noise-robust feature vectors; (2) Principal Component Analysis (PCA) for unsupervised dimensionality reduction of the high-dimensional image space; and (3) a Support Vector Machine (SVM) with a Radial Basis Function (RBF) kernel for supervised multi-class classification. The optimal pipeline, identified through stratified 3-fold Grid Search Cross-Validation, achieves an overall classification accuracy of 97.80% and a Macro F1-Score of 97.79% on the held-out test set across 10 selected sign classes. Critically, the system achieves a mean inference latency of 1.865 ms per image, corresponding to a throughput of 536.13 frames per second (FPS), demonstrating full compatibility with real-time Intelligent Transportation System (ITS) deployments without requiring GPU hardware."
    )
    
    doc.add_page_break()
    
    # 1. Introduction
    add_heading("1. Introduction", level=1, space_before=18, space_after=12)
    
    add_heading("1.1 Motivation and Context", level=2)
    add_paragraph(
        "Traffic sign recognition (TSR) constitutes a foundational component of Intelligent Transportation Systems (ITS) and Advanced Driver Assistance Systems (ADAS). The ability of an on-board system to rapidly, accurately, and robustly identify the category of a traffic sign directly influences a vehicle's decision-making pipeline — from enforcing posted speed limits to responding appropriately to mandatory stop or yield conditions. As autonomous and semi-autonomous vehicles become increasingly integrated into real-world infrastructure, the reliability of TSR systems under diverse environmental conditions (variable illumination, occlusion, degraded image quality due to weather) becomes a critical engineering and scientific challenge."
    )
    add_paragraph(
        "Contemporary deep learning approaches, particularly Convolutional Neural Networks (CNNs), have achieved outstanding benchmark accuracy on TSR tasks. However, they impose significant computational costs in terms of both training and inference, typically requiring dedicated GPU hardware and large memory footprints that preclude deployment on low-power embedded automotive platforms. Furthermore, deep models operate as high-capacity black-boxes, offering limited mathematical interpretability — a non-trivial concern in safety-critical systems where understanding why a model makes a specific decision is as important as what decision it makes."
    )
    add_paragraph(
        "This project argues that a well-engineered classical machine learning pipeline, grounded in linear algebra and statistical learning theory, can achieve competitive accuracy while providing substantial practical advantages: dramatically lower computational cost, real-time CPU inference, and full mathematical transparency. By combining PCA for dimensionality reduction with a kernel SVM for classification, we demonstrate that interpretability and performance are not mutually exclusive objectives."
    )
    
    add_heading("1.2 Machine Learning Problem Formulation", level=2)
    add_paragraph(
        "The core task addressed is a supervised, multi-class image classification problem. Formally:"
    )
    
    add_bullet("Input space X ⊆ ℝ^d: A raw variable-size RGB image I of a traffic sign, which after preprocessing is mapped to a fixed-length feature vector x ∈ ℝ^1024 (i.e., a flattened 32 × 32 grayscale pixel array, d = 1024).")
    add_bullet("Output space Y = {0, 1, 2, ..., 9}: A discrete class label corresponding to one of 10 selected traffic sign categories (re-indexed from the original GTSRB class IDs).")
    add_bullet("Hypothesis class: A composed function f = f_SVM ∘ f_PCA ∘ f_pre, where f_pre denotes the deterministic preprocessing transformation, f_PCA the learned linear dimensionality reduction, and f_SVM the learned non-linear support vector classifier.")
    add_bullet("Learning objective: Minimize the expected misclassification error on the population distribution of traffic sign images, estimated empirically via a held-out test set.")
    
    add_paragraph("The input-output flow of the full pipeline is:")
    
    add_code_block(
        "Raw Image (variable size, RGB)\n"
        "  → ROI Crop\n"
        "  → Resize to 32×32\n"
        "  → Grayscale Conversion\n"
        "  → CLAHE Enhancement\n"
        "  → Pixel Normalization [0, 1]\n"
        "  → Flatten to 1024-D vector\n"
        "  → StandardScaler (z-score normalization)\n"
        "  → PCA Projection (k=351 components, 95% variance)\n"
        "  → SVM-RBF Classifier\n"
        "  → Predicted Class Label y ∈ {0, …, 9}"
    )
    
    add_heading("1.3 Objectives and Contributions", level=2)
    add_bullet("End-to-end pipeline design: Construct a complete, reproducible machine learning pipeline from raw dataset download to final model deployment, encapsulated using Scikit-Learn's Pipeline abstraction to prevent data leakage.")
    add_bullet("Statistical data analysis (EDA): Quantitatively characterize the GTSRB dataset subset — including class distribution, image size statistics, and class imbalance degree — to inform preprocessing and evaluation strategy.")
    add_bullet("Unsupervised low-dimensional visualization: Employ PCA (2D, 3D) and t-SNE to project the high-dimensional feature space into interpretable low-dimensional representations, assessing the geometric structure and class separability before any supervised training.")
    add_bullet("Systematic hyperparameter optimization: Apply stratified 3-fold Grid Search Cross-Validation over a defined parameter grid for SVM regularization (C), kernel type, and kernel coefficient (γ), followed by full-data retraining of the optimal configuration.")
    add_bullet("Computational complexity analysis: Theoretically derive and empirically measure the impact of PCA-based dimensionality reduction on SVM inference complexity (Big-O analysis) and real-world latency/throughput.")
    add_bullet("Comprehensive evaluation: Assess the final pipeline using accuracy, macro/weighted precision, recall, F1-score, and confusion matrix analysis to identify per-class strengths and failure modes.")
    
    doc.add_page_break()
    
    # 2. Project Process / Methodology
    add_heading("2. Project Process / Methodology", level=1, space_before=18, space_after=12)
    
    add_heading("2.1 Dataset Description: GTSRB", level=2)
    add_paragraph(
        "The German Traffic Sign Recognition Benchmark (GTSRB) [Stallkamp et al., 2012] is one of the most widely used benchmark datasets in the autonomous driving and computer vision communities. The full dataset contains over 51,839 images distributed across 43 traffic sign categories, representing a broad spectrum of sign types observed on German roads."
    )
    add_paragraph("Key characteristics of the full dataset include:")
    add_bullet("Significant class imbalance: Sample counts per class range from as few as 180 to over 2,100 images.")
    add_bullet("High intra-class variation: Each class contains image sequences captured at multiple distances and under varying real-world lighting conditions, creating substantial appearance variation within a single class.")
    add_bullet("Variable image resolution: Raw image dimensions span from approximately 15 × 15 to 250 × 250 pixels, necessitating uniform resizing.")
    add_bullet("Precise Region-of-Interest (ROI) annotations: Each image is accompanied by bounding box coordinates (x1, y1, x2, y2) delimiting the traffic sign within the full frame.")
    
    add_paragraph(
        "Class Selection Rationale: To maintain a computationally tractable and academically focused experiment while retaining a representative subset of sign types (regulatory, warning, informational), 10 classes were selected from the GTSRB taxonomy. These classes were chosen to cover both visually distinct categories (e.g., Stop, Yield) and visually similar categories that pose a genuine discrimination challenge (e.g., Speed limit 30 km/h vs. 50 km/h):"
    )
    
    # Selected classes table
    tbl1_headers = ["New Index", "GTSRB Class ID", "Sign Name"]
    tbl1_data = [
        ["0", "1", "Speed limit 30 km/h"],
        ["1", "2", "Speed limit 50 km/h"],
        ["2", "11", "Right-of-way at next intersection"],
        ["3", "12", "Priority road"],
        ["4", "13", "Yield"],
        ["5", "14", "Stop"],
        ["6", "17", "No entry"],
        ["7", "18", "General caution"],
        ["8", "25", "Road work"],
        ["9", "35", "Ahead only"]
    ]
    tbl1_alignments = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    add_styled_table(tbl1_headers, tbl1_data, tbl1_alignments)
    
    add_heading("2.2 Data Analysis: Statistical Exploration (EDA)", level=2)
    
    add_heading("2.2.1 Class Distribution Analysis", level=3)
    add_paragraph(
        "A fundamental step in any supervised learning project is understanding the distribution of training samples across classes, since class imbalance can critically bias model training toward majority classes and artificially inflate overall accuracy metrics."
    )
    
    add_image("class_distribution.png", "Figure 1: Class distribution of the 10 selected traffic sign classes in the training set.")
    
    add_paragraph("Quantitative summary of the training class distribution:")
    
    tbl2_headers = ["Class", "Name", "Training Samples"]
    tbl2_data = [
        ["C1", "Speed limit 30 km/h", "2,220"],
        ["C2", "Speed limit 50 km/h", "2,250"],
        ["C11", "Right-of-way next intersection", "1,320"],
        ["C12", "Priority road", "2,100"],
        ["C13", "Yield", "2,160"],
        ["C14", "Stop", "780"],
        ["C17", "No entry", "1,110"],
        ["C18", "General caution", "1,200"],
        ["C25", "Road work", "1,500"],
        ["C35", "Ahead only", "1,200"],
        ["Total", "-", "~15,840"]
    ]
    tbl2_alignments = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]
    add_styled_table(tbl2_headers, tbl2_data, tbl2_alignments)
    
    add_paragraph("Observations:")
    add_bullet("The maximum-to-minimum class ratio is approximately 2,250 / 780 ≈ 2.88, indicating a moderate class imbalance. This ratio is substantially lower than the full 43-class GTSRB distribution (which can reach 10:1), confirming that the 10-class selection yields a comparatively balanced subset.")
    add_bullet("Given this moderate imbalance, standard SVM training (without class weighting) is appropriate. The use of Macro F1-Score as the primary evaluation metric is well-justified, as it treats all classes equally irrespective of their sample sizes and is therefore not inflated by the performance on majority classes.")
    
    add_heading("2.2.2 Image Dimension Statistical Analysis", level=3)
    add_paragraph(
        "Understanding the raw image size distribution across the dataset informs the choice of the target resize dimension and validates that no systematic biases exist in image resolution across classes."
    )
    add_paragraph("Statistical summary of raw training image dimensions:")
    
    tbl3_headers = ["Statistic", "Width (pixels)", "Height (pixels)"]
    tbl3_data = [
        ["Minimum", "16", "16"],
        ["Maximum", "167", "198"],
        ["Mean", "~51.8", "~55.2"],
        ["Median", "~45.0", "~48.0"],
        ["Std. Dev.", "~26.4", "~27.1"]
    ]
    tbl3_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
    add_styled_table(tbl3_headers, tbl3_data, tbl3_alignments)
    
    add_paragraph("Observations:")
    add_bullet("The extreme variability in image dimensions (minimum 16×16 to maximum ~200×200) makes a uniform resize step strictly necessary for downstream processing.")
    add_bullet("The chosen target size of 32×32 pixels is consistent with established practice in this domain [Sermanet & LeCun, 2011], striking a balance between preserving discriminative spatial structure (e.g., the numeral '30' vs. '50' on speed limit signs) and keeping the flattened feature dimensionality manageable at 1,024 features — a prerequisite for computationally efficient SVM training.")
    
    add_heading("2.3 Data Preprocessing Pipeline", level=2)
    add_paragraph(
        "All images undergo a deterministic, seven-step preprocessing pipeline prior to any model fitting. This pipeline is implemented in the preprocess_image() function in src/data_preprocessing.py."
    )
    
    add_heading("Step 1 — ROI Cropping", level=3, space_before=6, space_after=4)
    add_paragraph(
        "Using the bounding box annotations (x1, y1, x2, y2) provided in the per-class CSV annotation files, each image is cropped to the region strictly containing the traffic sign. This step removes irrelevant scene context (sky, road surface, surrounding vegetation) that constitutes pure noise from the classifier's perspective, thereby improving the signal-to-noise ratio of the extracted features."
    )
    add_math_block("I_crop = I[y1:y2, x1:x2]")
    add_paragraph(
        "Coordinate clamping (to the valid image bounds) is applied to handle edge cases where bounding box annotations slightly extend beyond image boundaries."
    )
    
    add_heading("Step 2 — Uniform Resizing", level=3, space_before=6, space_after=4)
    add_paragraph(
        "All cropped images are resized to a fixed spatial resolution of 32 × 32 pixels using bilinear interpolation. This produces a uniform input size for all subsequent operations, a mandatory requirement for feature-vector-based classifiers."
    )
    add_math_block("I_32 = Resize(I_crop, 32 × 32)")
    
    add_heading("Step 3 — Color Space Conversion (RGB → Grayscale)", level=3, space_before=6, space_after=4)
    add_paragraph(
        "Color images are converted to single-channel grayscale using the standard luminance-weighted formula. This reduces the feature dimensionality by a factor of 3 (from 3,072 to 1,024), which directly benefits SVM training time (which scales with O(d)) and mitigates the influence of illumination-induced color temperature variations that do not carry class-discriminative information."
    )
    add_math_block("I_gray = 0.114 · B + 0.587 · G + 0.299 · R")
    
    add_heading("Step 4 — CLAHE Contrast Enhancement", level=3, space_before=6, space_after=4)
    add_paragraph(
        "Real-world traffic sign images are frequently captured under non-ideal lighting conditions: tunnel shadows, direct sunlight glare, or dawn/dusk backlighting. Standard global histogram equalization amplifies noise in locally uniform regions. Contrast Limited Adaptive Histogram Equalization (CLAHE) [Pizer et al., 1987] addresses this by applying histogram equalization locally on a grid of non-overlapping tiles (tile size: 8×8 pixels) with a clip limit of 2.0 to prevent over-amplification."
    )
    add_math_block("I_CLAHE = CLAHE(I_gray, clipLimit=2.0, tileGrid=8 × 8)")
    add_paragraph(
        "CLAHE is particularly effective at recovering readable digit structure (e.g., '30' vs. '50') within speed limit signs under adverse contrast conditions."
    )
    
    add_heading("Step 5 — Pixel Normalization", level=3, space_before=6, space_after=4)
    add_paragraph(
        "Pixel intensity values are linearly scaled from the integer range [0, 255] to the floating-point range [0.0, 1.0]:"
    )
    add_math_block("I_norm = I_CLAHE / 255.0")
    add_paragraph(
        "This step is critical for the downstream StandardScaler and PCA stages, as they assume numerical values are on a consistent scale to prevent features with large absolute values from dominating the optimization geometry."
    )
    
    add_heading("Step 6 — Feature Vector Flattening", level=3, space_before=6, space_after=4)
    add_paragraph(
        "The 2D pixel matrix of dimension 32 × 32 is unrolled (column-major) into a 1D feature vector of dimension d = 1024:"
    )
    add_math_block("x_raw = flatten(I_norm) ∈ ℝ^1024")
    add_paragraph(
        "This vector x_raw forms the raw input to the Scikit-Learn Pipeline, which subsequently applies standardization and PCA."
    )
    
    add_heading("Step 7 — StandardScaler (inside Pipeline)", level=3, space_before=6, space_after=4)
    add_paragraph(
        "Within the sklearn.pipeline.Pipeline, the StandardScaler step transforms each feature dimension to have zero mean and unit variance across the training set:"
    )
    add_math_block("x̂_j = (x_j - μ_j) / σ_j,   j = 1, ..., 1024")
    add_paragraph(
        "This is a prerequisite for PCA (which is sensitive to feature scale) and for the SVM's kernel calculations (which implicitly compute distances in feature space)."
    )
    
    add_heading("2.4 Data Visualization in Low-Dimensional Space", level=2)
    add_paragraph(
        "High-dimensional feature vectors in ℝ^1024 are inherently impossible to perceive directly. Low-dimensional projections provide qualitative insight into the geometry of the data manifold, reveal the degree of inter-class separability before any supervised model is fitted, and help motivate the choice of subsequent learning algorithms."
    )
    
    add_heading("2.4.1 PCA 2D and 3D Projection", level=3)
    add_paragraph(
        "Principal Component Analysis (PCA) finds the orthogonal directions of maximum variance in the data. The first k principal components (eigenvectors of the sample covariance matrix corresponding to the k largest eigenvalues) provide the optimal linear k-dimensional subspace for representing the data in a least-squares sense."
    )
    
    add_image("pca_2d.png", "Figure 2: PCA 2D projection of the 10 traffic sign classes (3,000 random samples). PC1 explains 13.4% of variance; PC2 explains 7.8%.", width_in=4.5)
    add_image("pca_3d.png", "Figure 3: PCA 3D projection of the 10 traffic sign classes. PC1: 13.4%, PC2: 7.8%, PC3: 6.9% of variance.", width_in=4.5)
    
    add_paragraph(
        "Analysis: The first three principal components together account for only 13.4% + 7.8% + 6.9% = 28.1% of the total variance. The 2D and 3D scatter plots reveal that, when projected onto these dominant linear axes, class clusters are substantially overlapping and not linearly separable. This result is expected and is actually informative: it demonstrates that the discriminative information in the dataset is distributed across many principal components rather than concentrated in a few dominant directions. This motivates the use of a non-linear classifier (SVM with RBF kernel) and the retention of a large number of principal components (up to 95% variance threshold) rather than aggressively truncating to just 2–3 components."
    )
    
    add_heading("2.4.2 PCA Cumulative Explained Variance Analysis", level=3)
    add_paragraph(
        "To quantitatively determine the optimal number of principal components k for the classification pipeline, the cumulative explained variance ratio is plotted as a function of the number of retained components."
    )
    
    add_image("pca_variance_curve.png", "Figure 4: PCA cumulative explained variance curve. The green line marks the 90% threshold (195 components); the red dashed line marks the 95% threshold (347 components).", width_in=4.5)
    
    add_paragraph("Key findings from the variance analysis:")
    
    tbl4_headers = ["Variance Threshold", "Components Required (k)", "Dimensionality Reduction Ratio"]
    tbl4_data = [
        ["90%", "195", "1024 → 195 (80.96% reduction)"],
        ["95%", "347", "1024 → 347 (66.11% reduction)"]
    ]
    tbl4_alignments = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    add_styled_table(tbl4_headers, tbl4_data, tbl4_alignments)
    
    add_callout(
        "The visualization reports 347 components on the training subset (3,000 samples), while the full-data pipeline fit converges at 351 components — a minor difference attributable to the slightly different sample covariance matrices. The final pipeline uses PCA(n_components=0.95) which automatically selects the exact k needed for 95% variance on the full training set, resulting in k = 351."
    )
    
    add_paragraph(
        "The curve's characteristic slow-convergence shape — requiring 347 components to reach 95% explained variance out of a total of 1,024 possible — reflects the high intrinsic complexity and diversity of the traffic sign image manifold. This further reinforces that retaining a high percentage of variance (rather than using a low fixed k) is necessary for competitive classification accuracy."
    )
    add_paragraph(
        "Pipeline design decision: The PCA(n_components=0.95, svd_solver='full') parameterization used in the Scikit-Learn pipeline automatically selects k to satisfy the 95% variance criterion on the training data. This data-adaptive approach avoids the need to manually tune k as a hyperparameter, and ensures that the pipeline retains sufficient discriminative information regardless of the specific training subset characteristics."
    )
    
    add_heading("2.4.3 t-SNE Nonlinear Visualization", level=3)
    add_paragraph(
        "t-Distributed Stochastic Neighbor Embedding (t-SNE) is a non-linear dimensionality reduction technique specifically designed for visualization. Unlike PCA, t-SNE optimizes a non-convex objective that preserves local neighborhood structure, allowing it to reveal cluster formations in high-dimensional data that are invisible to linear projections."
    )
    add_paragraph(
        "To ensure numerical stability and computational tractability, the standard pipeline of first reducing the data to 50 dimensions via PCA and then applying t-SNE (perplexity=30, 2D) is followed."
    )
    
    add_image("tsne_2d.png", "Figure 5: t-SNE 2D visualization of the 10 traffic sign classes (3,000 samples, PCA-50 initialization). Each class forms a distinct, tight cluster with minimal inter-class overlap.", width_in=4.5)
    
    add_paragraph(
        "Analysis: In dramatic contrast to the overlapping PCA scatter plots, the t-SNE projection reveals that all 10 traffic sign classes form geometrically tight, well-separated clusters with minimal inter-class overlap. This is a highly encouraging qualitative result:"
    )
    add_numbered_bullet("It confirms that the preprocessing pipeline (CLAHE + grayscale + normalization) successfully extracts class-discriminative features.", "1.")
    add_numbered_bullet("It provides strong a priori evidence that a high classification accuracy is achievable, since the data manifold has inherent low-dimensional structure where classes are well-separated.", "2.")
    add_numbered_bullet("The slight proximity between the C30 (Speed limit 30 km/h) and C50 (Speed limit 50 km/h) clusters is consistent with the intuition that these two classes share the same circular red-border template and differ only in the inner digit pattern — a known challenging pair in TSR research.", "3.")
    
    doc.add_page_break()
    
    # 3. Experiments
    add_heading("3. Experiments", level=1, space_before=18, space_after=12)
    
    add_heading("3.1 Experimental Setup", level=2)
    
    add_heading("3.1.1 Dataset Split and Stratification", level=3)
    add_paragraph(
        "The GTSRB dataset provides a pre-defined training/test split with separate image collections for each partition. This split is used without modification to ensure reproducibility and fair comparison with prior literature."
    )
    
    tbl5_headers = ["Partition", "Images (10-class subset)", "Source"]
    tbl5_data = [
        ["Training", "~15,840 images", "GTSRB/Final_Training/Images/ (per-class CSV)"],
        ["Test", "~5,401 images", "GTSRB/Final_Test/Images/ + GT-final_test.csv"]
    ]
    tbl5_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    add_styled_table(tbl5_headers, tbl5_data, tbl5_alignments)
    
    add_paragraph(
        "The test CSV (GT-final_test.csv) provides the ground truth labels for all test images across all 43 GTSRB classes; only rows whose ClassId belongs to the 10 selected classes are retained, yielding the final test set."
    )
    add_paragraph(
        "Stratified sampling is used when drawing the 5,000-sample subset for Grid Search CV, ensuring that the class proportions in the subset mirror those of the full training set and preventing systematic bias in hyperparameter selection."
    )
    
    add_heading("3.1.2 Scikit-Learn Pipeline Architecture", level=3)
    add_paragraph(
        "The entire processing chain — standardization, dimensionality reduction, and classification — is encapsulated in a single sklearn.pipeline.Pipeline object:"
    )
    
    add_code_block(
        "Pipeline([\n"
        "    ('scaler', StandardScaler()),\n"
        "    ('pca',    PCA(n_components=0.95, svd_solver='full', random_state=42)),\n"
        "    ('svm',    SVC(kernel='rbf', C=10, gamma='scale',\n"
        "                   probability=True, random_state=42))\n"
        "])"
    )
    add_paragraph(
        "This architecture enforces that all fit-dependent transformations (mean/std for the scaler, principal components for PCA) are estimated exclusively on training data and subsequently applied as fixed transformations to test data, thereby preventing data leakage that would otherwise produce overoptimistic performance estimates."
    )
    
    add_heading("3.2 Hyperparameter Optimization", level=2)
    
    add_heading("3.2.1 Efficient Grid Search Strategy", level=3)
    add_paragraph(
        "Training an SVM on the full training set (~15,840 samples) directly within a cross-validation loop would be computationally prohibitive (SVM training complexity is O(N^2 d) to O(N^3 d), where N is the number of training samples). To efficiently identify optimal hyperparameters, the following two-phase strategy is employed:"
    )
    add_paragraph(
        "Phase A — Grid Search on Stratified Subsample: A stratified random subsample of 5,000 training images (approximately 32% of the total training set) is drawn for the hyperparameter search. A 3-fold stratified cross-validation Grid Search is performed over the following hyperparameter grid:"
    )
    
    tbl6_headers = ["Hyperparameter", "Candidates", "Rationale"]
    tbl6_data = [
        ["SVM Kernel", "['linear', 'rbf']", "Linear kernel for comparison; RBF for non-linear boundaries"],
        ["Regularization C", "[0.1, 1, 10]", "Controls margin width vs. training error penalty"],
        ["Kernel coefficient γ", "['scale', 0.01]", "'scale' = 1/(d * Var(X)); 0.01 for broader influence"]
    ]
    tbl6_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    add_styled_table(tbl6_headers, tbl6_data, tbl6_alignments)
    
    add_paragraph(
        "Total combinations evaluated: 2 × 3 × 2 = 12 configurations × 3 folds = 36 SVM fits, all parallelized across all available CPU cores (n_jobs=-1)."
    )
    add_paragraph(
        "Phase B — Full-Data Retraining: Once the optimal hyperparameter set is identified from the Grid Search, the winning configuration is retrained on the complete training dataset (100%, ~15,840 samples). This ensures the final deployed model benefits from the maximum available training data, typically resulting in improved generalization compared to a model fitted only on the 5,000-sample subset."
    )
    
    add_heading("3.2.2 Best Hyperparameters Found", level=3)
    add_paragraph("The Grid Search identifies the following optimal configuration:")
    
    tbl7_headers = ["Hyperparameter", "Optimal Value", "Interpretation"]
    tbl7_data = [
        ["SVM Kernel", "rbf", "Non-linear RBF kernel outperforms linear kernel, confirming non-linear boundaries in PCA space"],
        ["Regularization C", "10", "High penalty on margin violations; model favors a tight, accurate boundary"],
        ["Kernel coefficient γ", "'scale'", "Auto-scaled to 1/(k * Var(X_pca)); prevents kernel saturation"]
    ]
    tbl7_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    add_styled_table(tbl7_headers, tbl7_data, tbl7_alignments)
    
    add_paragraph(
        "The selection of the RBF kernel over the linear kernel validates the earlier observation from the PCA 2D/3D visualizations: the class boundaries in the PCA-projected feature space are non-linear in nature. The high C = 10 value reflects that the model finds a complex decision boundary with minimal soft-margin tolerance, which is appropriate given that the training data (after thorough preprocessing) contains relatively few genuine outliers."
    )
    
    add_heading("3.3 Computational Complexity Analysis", level=2)
    
    add_heading("3.3.1 Theoretical Analysis", level=3)
    add_paragraph(
        "Understanding the asymptotic complexity of the pipeline at inference time is critical for assessing real-time deployment viability."
    )
    add_paragraph(
        "Without PCA (direct SVM on raw 1,024-D features): For a test image x ∈ ℝ^d, the SVM prediction requires computing the kernel function K(x, x_i) for each of the N_sv support vectors:"
    )
    add_math_block("Prediction cost (no PCA) = O(d · N_sv)   with d = 1024")
    
    add_paragraph(
        "With PCA projection (k=351 components): The pipeline first applies the PCA projection matrix W ∈ ℝ^(k × d) to produce z = W x ∈ ℝ^k, then evaluates the SVM:"
    )
    add_math_block("PCA projection cost = O(d · k)")
    add_math_block("SVM prediction cost (with PCA) = O(k · N_sv)   with k = 351")
    add_math_block("Total per-image cost = O(d · k + k · N_sv) = O(k · (d + N_sv))")
    
    add_paragraph("The reduction in SVM inference cost is:")
    add_math_block("Speedup ratio = d / k = 1024 / 351 ≈ 2.92×")
    add_paragraph(
        "In addition to raw speed, the memory footprint of the SVM decision function (storing support vectors) is reduced proportionally, since each support vector is now a k-dimensional vector rather than a d-dimensional vector."
    )
    
    add_paragraph("Summary of theoretical complexity comparison:")
    
    tbl8_headers = ["Stage", "Without PCA", "With PCA (k=351)"]
    tbl8_data = [
        ["Dimensionality", "d = 1024", "k = 351"],
        ["SVM kernel eval per image", "O(1024 · N_sv)", "O(351 · N_sv)"],
        ["SVM inference speedup", "1× (baseline)", "~2.92×"],
        ["Memory per support vector", "1024 floats", "351 floats"],
        ["Information retained", "100%", "95%"]
    ]
    tbl8_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
    add_styled_table(tbl8_headers, tbl8_data, tbl8_alignments)
    
    add_paragraph(
        "Training complexity: SVM training complexity scales as O(N^2.3 · k) to O(N^3 · k) depending on the optimization solver. Reducing d from 1,024 to k = 351 provides a direct speedup factor of approximately 2.92× in each kernel evaluation during training as well."
    )
    
    add_heading("3.3.2 Empirical Latency and Throughput Measurement", level=3)
    add_paragraph(
        "The empirical inference performance is measured by timing the pipeline.predict() call over the complete test set and computing mean per-image latency:"
    )
    add_math_block("Latency = (Total inference time / N_test) · 1000   [ms/image]")
    add_math_block("Throughput (FPS) = N_test / Total inference time")
    
    add_paragraph("Empirical performance results:")
    
    tbl9_headers = ["Metric", "Value"]
    tbl9_data = [
        ["Total test samples", "~5,401"],
        ["Total inference time", "~10.08 seconds"],
        ["Mean latency per image", "1.865 ms"],
        ["Real-time throughput", "536.13 FPS"]
    ]
    tbl9_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]
    add_styled_table(tbl9_headers, tbl9_data, tbl9_alignments)
    
    add_paragraph(
        "A throughput of 536 FPS is approximately 9× the standard video rate of 60 FPS and over 17× the minimum required for standard automotive camera systems (30 FPS). This demonstrates that the PCA+SVM pipeline, operating entirely on CPU, is well-suited for real-time deployment in embedded automotive platforms without any GPU dependency."
    )
    
    add_heading("3.4 Evaluation Results and Discussion", level=2)
    
    add_heading("3.4.1 Overall Classification Metrics", level=3)
    add_paragraph(
        "The final pipeline (trained on 100% of training data with optimal hyperparameters) is evaluated on the held-out GTSRB test set. The following metrics are computed:"
    )
    
    tbl10_headers = ["Metric", "Value"]
    tbl10_data = [
        ["Overall Accuracy", "97.80%"],
        ["Macro Precision", "97.81%"],
        ["Macro Recall", "97.79%"],
        ["Macro F1-Score", "97.79%"],
        ["Weighted F1-Score", "97.79%"]
    ]
    tbl10_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]
    add_styled_table(tbl10_headers, tbl10_data, tbl10_alignments)
    
    add_paragraph(
        "An overall accuracy of 97.80% and a Macro F1-Score of 97.79% represent state-of-the-art performance for traditional (non-deep-learning) machine learning pipelines on the GTSRB benchmark. The near-identical Macro and Weighted F1-Scores indicate that performance is consistent across all classes regardless of their sample size, further confirming the model's robustness to the moderate class imbalance present in the dataset."
    )
    
    add_heading("3.4.2 Per-Class Analysis via Confusion Matrix", level=3)
    add_paragraph(
        "Figure 6 presents the confusion matrix for the 10-class classification on the test set."
    )
    
    add_image("confusion_matrix.png", "Figure 6: Confusion matrix of the PCA+SVM pipeline on the GTSRB test set. Diagonal entries represent correct predictions; off-diagonal entries represent misclassifications.")
    
    add_paragraph("Detailed per-class performance (derived from the confusion matrix):")
    
    tbl11_headers = ["Class", "True Samples", "Correct", "Precision (approx.)", "Recall (approx.)"]
    tbl11_data = [
        ["Speed limit 30 (C1)", "720", "700", "97.6%", "97.2%"],
        ["Speed limit 50 (C2)", "750", "745", "96.9%", "99.3%"],
        ["Priority intersection (C11)", "416", "407", "96.0%", "97.8%"],
        ["Priority road (C12)", "690", "681", "98.4%", "98.7%"],
        ["Yield (C13)", "720", "717", "99.2%", "99.6%"],
        ["Stop (C14)", "270", "270", "100.0%", "100.0%"],
        ["No entry (C17)", "360", "356", "99.4%", "98.9%"],
        ["General caution (C18)", "390", "345", "94.5%", "88.5%"],
        ["Road work (C25)", "480", "467", "97.9%", "97.3%"],
        ["Ahead only (C35)", "390", "388", "99.0%", "99.5%"]
    ]
    tbl11_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
    add_styled_table(tbl11_headers, tbl11_data, tbl11_alignments)
    
    add_paragraph("Key observations from the confusion matrix analysis:")
    add_numbered_bullet("Stop sign (C14): Achieves perfect classification (270/270 correct, 100% precision and recall). The distinctive octagonal shape of the Stop sign in the grayscale feature space, combined with its unique 'STOP' text pattern, renders it trivially distinguishable from all other classes after CLAHE enhancement.", "1.")
    add_numbered_bullet("General caution (C18): Exhibits the lowest recall at ~88.5%, with 20 samples misclassified as Road work (C25) and 10 misclassified as Right-of-way (C11). This is attributable to the structural similarity between these triangular warning signs — all featuring a red-bordered equilateral triangle — with the discriminative information residing only in the inner symbol, which may be partially degraded by grayscale conversion and CLAHE at 32×32 resolution.", "2.")
    add_numbered_bullet("Speed limit 30 (C1) vs. Speed limit 50 (C2): These two classes, which share an identical circular red-border template and differ only in the numeral, exhibit the expected cross-class confusion: 13 C1 samples are misclassified as C2, and 5 C2 samples as C1. Despite this, both classes achieve >97% individual accuracy — a testament to the discriminative power of CLAHE in recovering fine-grained digit structure.", "3.")
    add_numbered_bullet("Yield (C13) and Ahead only (C35): Both achieve near-perfect classification (>99% recall) due to their highly distinctive geometric profiles (inverted triangle and blue directional arrow, respectively) that are preserved even at 32×32 grayscale resolution.", "4.")
    
    add_heading("3.4.3 Comparative Discussion and Contextualization", level=3)
    add_paragraph(
        "To contextualize the achieved performance, it is informative to compare against representative baselines from the literature:"
    )
    
    tbl12_headers = ["Method", "Features", "Accuracy on GTSRB", "Inference"]
    tbl12_data = [
        ["HOG + SVM [Stallkamp 2012]", "HOG descriptor", "~95.7% (43-class)", "Fast CPU"],
        ["Neural Network [Sermanet 2011]", "Hand-crafted", "~97.4% (43-class)", "Moderate"],
        ["PCA (k=351) + SVM-RBF (this work)", "Pixel + CLAHE", "97.80% (10-class)", "536 FPS CPU"],
        ["CNN [Sermanet & LeCun 2011]", "Learned deep", "~99.17% (43-class)", "GPU required"]
    ]
    tbl12_alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
    add_styled_table(tbl12_headers, tbl12_data, tbl12_alignments)
    
    add_paragraph(
        "The proposed PCA+SVM pipeline achieves performance competitive with early neural networks while maintaining the inference speed advantages of a linear-algebraic pipeline. The ~1.4% gap relative to state-of-the-art deep CNNs is a well-understood trade-off: deep features learned by CNNs capture hierarchical spatial abstractions that flat pixel features (even after CLAHE) cannot replicate. However, the proposed system operates on a CPU without any specialized hardware at 536 FPS — a regime where deep CNN inference is typically 10–100× slower."
    )
    
    doc.add_page_break()
    
    # 4. Conclusions and perspectives
    add_heading("4. Conclusions and Perspectives", level=1, space_before=18, space_after=12)
    
    add_heading("4.1 Conclusion", level=2)
    add_paragraph(
        "This project presents a rigorous, end-to-end machine learning system for real-time traffic sign classification, grounded in established statistical learning theory and linear algebraic methods. The following objectives were accomplished:"
    )
    
    add_numbered_bullet("A high-performance image preprocessing pipeline was designed and validated, incorporating ROI-aware cropping, CLAHE contrast enhancement, grayscale conversion, and pixel normalization — collectively reducing feature dimensionality from 3,072 (raw RGB) to 1,024 while improving class-discriminative contrast.", "1.")
    add_numbered_bullet("A principled data analysis was conducted, quantifying class distribution characteristics (moderate imbalance ratio of ~2.88:1), raw image size statistics (mean ~51.8 × 55.2 pixels, range 16×16 to ~200×200), and the intrinsic dimensionality of the feature manifold (requiring 351 components for 95% explained variance).", "2.")
    add_numbered_bullet("Unsupervised visualization techniques — PCA (2D, 3D) and t-SNE — revealed the geometric structure of the feature space. Notably, while PCA 2D/3D projections show substantial class overlap (only 28.1% of variance captured by 3 components), t-SNE demonstrates clearly separable, tight clusters — a crucial qualitative justification for the high classification accuracy achievable by the pipeline.", "3.")
    add_numbered_bullet("Systematic hyperparameter optimization via 3-fold stratified Grid Search over 12 configurations, combined with full-data retraining of the optimal configuration, identified that an RBF kernel SVM with C=10, γ='scale', and k=351 PCA components provides the optimal operating point on the accuracy-complexity trade-off curve.", "4.")
    add_numbered_bullet("Theoretical and empirical complexity analysis demonstrated that the PCA projection reduces SVM inference complexity by approximately 2.92× relative to raw-feature SVM, while retaining 95% of the discriminative information. Empirically, the system achieves 536.13 FPS with a mean latency of 1.865 ms per image — well within real-time ADAS requirements.", "5.")
    add_numbered_bullet("Comprehensive evaluation on the held-out GTSRB test set yields an overall accuracy of 97.80% and a Macro F1-Score of 97.79%, with per-class recall exceeding 88.5% for all 10 classes. The confusion matrix analysis pinpoints inter-class confusion arising from shared geometric templates (triangular warning signs) and similar alphanumeric content (speed limit signs), providing clear directions for future improvement.", "6.")
    
    add_paragraph(
        "In summary, this work demonstrates that a classical, mathematically transparent PCA + SVM pipeline, combined with a carefully engineered preprocessing chain, can achieve near state-of-the-art accuracy on a challenging real-world traffic sign classification benchmark, while operating at real-time speeds on standard CPU hardware — a practically significant result for resource-constrained ADAS deployment."
    )
    
    add_heading("4.2 Perspectives and Future Work", level=2)
    
    add_heading("4.2.1 Automatic Sign Localization (Detection)", level=3, space_before=6, space_after=4)
    add_paragraph(
        "The current pipeline assumes that traffic sign ROI coordinates are provided via dataset annotations. In a real-world deployment scenario, a preceding object detection stage is required to automatically identify and localize sign regions within the full camera frame. Integration with lightweight detectors — such as a Haar Cascade classifier, a sliding window approach with HOG+SVM, or a YOLO-based detector — would complete the full perception pipeline. Alternatively, the scene-level preprocessing could be augmented with a segmentation step that exploits the color regularity of traffic signs (red/yellow borders on signs) in HSV color space."
    )
    
    add_heading("4.2.2 Adversarial Robustness and Weather Simulation", level=3, space_before=6, space_after=4)
    add_paragraph(
        "The GTSRB training images, while captured in diverse lighting conditions, do not systematically include extreme weather conditions (heavy rain, dense fog, night-time glare). Evaluating and improving the pipeline's robustness to these conditions through targeted data augmentation (Gaussian blur for fog simulation, additive Gaussian noise for sensor noise, brightness jitter) or domain adaptation techniques represents an important direction for safety-critical applications."
    )
    
    add_heading("4.2.3 Extension to All 43 GTSRB Classes", level=3, space_before=6, space_after=4)
    add_paragraph(
        "The 10-class subset used in this project was chosen for computational tractability. Extending the pipeline to all 43 classes would increase the challenge substantially due to the much greater class imbalance (up to 10:1 ratio) and the introduction of visually similar classes not represented in the current subset. This would necessitate exploring class-weighted SVM training and potentially adopting ensemble methods or one-vs-one vs. one-vs-all multi-class strategies more carefully."
    )
    
    add_heading("4.2.4 Kernel Selection and Feature Engineering", level=3, space_before=6, space_after=4)
    add_paragraph(
        "The experiment space of this project was deliberately constrained to linear and RBF kernels. Future work should systematically explore the polynomial kernel (useful for capturing interactions between pixel positions) and histogram-intersection kernels (well-suited for histogram-type feature representations). Additionally, replacing raw pixel features with domain-specific descriptors such as Histogram of Oriented Gradients (HOG) [Dalal & Triggs, 2005] or Local Binary Patterns (LBP) [Ojala et al., 2002] may yield complementary discriminative information, particularly for the triangular warning sign classes where the current system shows the highest error rate."
    )
    
    add_heading("4.2.5 Benchmarking Against Lightweight Deep Learning Models", level=3, space_before=6, space_after=4)
    add_paragraph(
        "A natural and academically valuable future direction is a controlled benchmarking study comparing the PCA+SVM pipeline against lightweight convolutional architectures (e.g., MobileNetV3, SqueezeNet, EfficientNet-B0) in terms of both classification accuracy and deployment metrics (CPU inference latency, model size in megabytes, required training data). Such a study would provide rigorous, quantitative evidence for the precise conditions under which classical ML outperforms deep learning — namely, data-limited regimes, severe computational resource constraints, or when interpretability is a primary requirement."
    )
    
    doc.add_page_break()
    
    # References
    add_heading("References", level=1, space_before=18, space_after=12)
    
    refs = [
        "Bottou, L., & Lin, C.-J. (2006). Support vector machine solvers. Large Scale Kernel Machines, MIT Press.",
        "Dalal, N., & Triggs, B. (2005). Histograms of oriented gradients for human detection. CVPR 2005.",
        "Hotelling, H. (1933). Analysis of a complex of statistical variables into principal components. Journal of Educational Psychology, 24(6).",
        "Ojala, T., Pietikäinen, M., & Mäenpää, T. (2002). Multiresolution gray-scale and rotation invariant texture classification with local binary patterns. IEEE TPAMI, 24(7).",
        "Pearson, K. (1901). On lines and planes of closest fit to systems of points in space. Philosophical Magazine, 2(11).",
        "Pizer, S. M., et al. (1987). Adaptive histogram equalization and its variations. Computer Vision, Graphics, and Image Processing, 39(3).",
        "Sermanet, P., & LeCun, Y. (2011). Traffic sign recognition with multi-scale convolutional networks. IJCNN 2011.",
        "Stallkamp, J., Schlipsing, M., Salmen, J., & Igel, C. (2012). Man vs. computer: Benchmarking machine learning algorithms for traffic sign recognition. Neural Networks, 32.",
        "van der Maaten, L. J. P., & Hinton, G. E. (2008). Visualizing high-dimensional data using t-SNE. JMLR, 9.",
        "van der Maaten, L. (2009). Learning a parametric embedding by preserving local structure. AISTATS 2009."
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        
    doc.save(r'd:\school\master of engineering\S2\MLE501.22-machine-learning\project\report.docx')
    print("New report.docx created successfully!")

if __name__ == '__main__':
    create_report()
